# -*- coding：utf-8 -*-
# &Author  AnFany

#  生成数据集报告的主程序

from docx import Document  # 写入.word
from docx.shared import Inches  #
import io
import pandas as pd
import numpy as np
import auxiliary_fuction as a_f  # 辅助函数
import data_report_config as drc  # 数据报告的配置
import read_data as data  # 数据文件
from mpl_toolkits.mplot3d import Axes3D

# WORD版数据报告
DOC = Document()
DOC.add_heading('%s:数据分析报告' % drc.DATA_NAME, 0)  # word文档的标题


class VIEW():

    def __init__(self):
        """
        生成.word版的数据报告
        """
        # 需要进行分析的数据，DataFrame格式
        self.data = data.DataSet
        # 数据集中的目标字段
        self.name = drc.TARGET_NAME
        # 值的类型为数值型，但是看作离散性特征的字段名称的列表
        self.f_list = drc.TYPE_LIST
        # 存储图片的路径
        self.path = drc.SAVE_PATH
        # 数据报告的名称
        self.report_name = drc.DATA_NAME
        # 类别型特征中值的个数超过阈值就不再进行展示,需要进行特殊处理
        self.type_count = drc.THRESHOLD_VALUE
        # 缺失值的展示标识
        self.miss_name = '缺失值'
        # 类别型特征的值较多的字段
        self.special = []
        # 缺失值字段的字典
        self.miss = self.overview()
        # 只有一个值的特征字段
        self.one = self.distribution_feature()

    def overview(self):
        """
        数据概览，以及缺失值的信息
        :return:
        """
        miss_dict = {}
        for column_name in self.data.columns:
            length = self.data[column_name].count()
            if length != len(self.data):
                miss_dict[column_name] = '缺失率：%.3f' % (1 - length / len(self.data))

        DOC.add_heading('一、数据概览', level=1)
        buffer = io.StringIO()
        self.data.info(buf=buffer)
        s = buffer.getvalue()
        DOC.add_paragraph('%s' % s)
        print('一、数据概览：完毕')

        DOC.add_heading('二、缺失值分析', level=1)
        if miss_dict:
            DOC.add_paragraph('含有缺失值的特征：\n%s' % sorted(miss_dict.items(), key=lambda x: x[1], reverse=True))
        else:
            DOC.add_paragraph('所有特征均没有缺失值')
        print('二、缺失值分析：完毕')
        return miss_dict

    def distribution_feature(self):
        """
        单一特征的数据分布，数值型采用直方图，类别型用柱状图，
        对于数值型特征，删除缺失值。对于类别型特征，缺失值看作一类
        只有一个值的字段，无论是数值型还是类别型，均不绘制图
        :return: 以字段名命名的图
        """
        DOC.add_heading('三、每个字段的值的分布', level=1)
        # 只有一个值的字段，无论字段的属性是什么，均不绘制图
        one_number_list = []
        for key in self.data:
            # 类别型特征
            if self.data[key].dtype == object or key in self.f_list:
                plot_data = self.data[key]
                if len(set(list(plot_data.values))) != 1:
                    # 缺失值变为自定义的类别
                    plot_data = plot_data.replace(np.nan, self.miss_name).values
                    # 判断该特征的不同值的个数
                    length = len(set(plot_data))
                    if length > self.type_count:
                        self.special.append(key)
                        continue
                    # 产生数据
                    label_data, num_data = a_f.generate_bar(plot_data)
                    if key in self.miss:
                        # 获取图片的标题
                        if key == self.name:
                            title = '目标字段' + key + ', %s' % self.miss[key]
                        else:
                            title = '特征字段' + key + ', %s' % self.miss[key]
                    else:
                        # 绘制柱状图
                        if key == self.name:
                            title = '目标字段' + key
                        else:
                            title = '特征字段' + key
                    # 绘制柱状图
                    a_f.plot_bar(label_data, num_data, title, key)
                    # 输出到数据报告
                    DOC.add_paragraph('特征字段：' + key, style='List Number')
                    DOC.add_picture(r'%s/%s.png' % (self.path, key), width=Inches(5.8))
                else:
                    one_number_list.append(key)
                    DOC.add_paragraph('字段%s只有一个值' % key, style='List Number')

            else:
                # 绘制直方图
                plot_data = self.data[key].dropna().values  # 删除缺失值
                if len(set(list(plot_data))) != 1:
                    if key in self.miss:
                        if key == self.name:
                            title = '目标字段' + key + ', %s' % self.miss[key]
                        else:
                            title = '特征字段' + key + ', %s' % self.miss[key]
                    else:
                        if key == self.name:
                            title = '目标字段' + key
                        else:
                            title = '特征字段' + key
                    # 绘制直方图
                    a_f.plot_hist(plot_data, title, key)
                    # 输出到数据报告
                    DOC.add_paragraph('字段：' + key, style='List Number')
                    DOC.add_picture(r'%s/%s.png' % (self.path, key), width=Inches(5.8))
                else:
                    one_number_list.append(key)
                    DOC.add_paragraph('字段%s只有一个值' % key, style='List Number')
        print('三、单个字段的值的分布：完毕')
        return one_number_list

    def relation_feature_with_target(self):
        """
        获取每个特征字段与目标字段之间的关系：
        对于数值型特征，删除缺失值。对于类别型特征，缺失值看作一类
        如果目标字段是类别型特征，特征是类别型的，用分类柱状图；特征是数值的，用概率密度图
        如果目标字段是数值型特征，特征是类别型的，用概率密度图；特征是数值的，用散点图
        只有一个值的字段不绘制图
        :return: 绘制的图
        """
        DOC.add_heading('四、每个特征字段与目标字段之间的关系', level=1)
        for key in self.data:
            if key != self.name:
                if key not in self.one:
                    # 目标字段的值为类别型
                    if self.data[self.name].dtype == object or self.name in self.f_list:
                        # 特征字段值为类别型
                        if self.data[key].dtype == object or key in self.f_list:
                            # 类别的值
                            xdata = list(self.data[self.name].values)
                            if key in self.miss:
                                ydata = self.data[key].fillna('缺失值')
                            else:
                                ydata = self.data[key]

                            # 判断该类型中不同值的个数
                            if len(set(ydata.values)) > self.type_count:
                                continue
                            # 生成数据
                            x_data, y_data, data_dict = a_f.generate_bar_type(xdata, ydata)

                            if key not in self.miss:
                                title = "字段%s：不同类别的占比" % key
                            else:
                                title = "字段%s：不同类别的占比" % key + ' %s' % self.miss[key]
                            # 绘制分类柱状图
                            a_f.plot_hist_with_type(x_data, y_data, data_dict, title, key + '_' + self.name)

                        # 特征字段值为数值型
                        else:
                            new_df = pd.DataFrame()
                            new_df['type'] = self.data[self.name]
                            new_df['num'] = self.data[key]
                            new_df = new_df.dropna(axis=0, how='any')
                            x_data = new_df['type'].values
                            y_data = new_df['num'].values
                            data_dict = a_f.generate_data_for_plot_distribution(x_data, y_data)
                            # 获取标题
                            if key not in self.miss:
                                title = "字段%s：不同类别的分布" % key
                            else:
                                title = "字段%s：不同类别的分布" % key + ' %s' % self.miss[key]
                            # 绘制概率密度图
                            a_f.plot_density_with_type(data_dict, title, key + '_' + self.name, key)
                        # 输出数据报告
                        DOC.add_paragraph('关系: ' + key + ' VS ' + self.name, style='List Bullet')
                        DOC.add_picture(r'%s/%s_%s.png' % (self.path, key, self.name), width=Inches(5.8))
                    # 目标字段的值为数值型
                    else:
                        # 特征字段值为类别型
                        if self.data[key].dtype == object or key in self.f_list:
                            # 获取数据
                            new_df = pd.DataFrame()
                            new_df['num'] = self.data[self.name]
                            new_df['type'] = self.data[key]
                            new_df = new_df.dropna(axis=0, how='any')
                            x_data = new_df['type'].values
                            y_data = new_df['num'].values
                            # 判断该类型中不同值的个数
                            if len(set(y_data)) > self.type_count:
                                continue
                            data_dict = a_f.generate_data_for_plot_distribution(x_data, y_data)
                            # 获取标题
                            if key not in self.miss:
                                title = "特征%s：分布" % key
                            else:
                                title = "特征%s：分布" % key + ' %s' % self.miss[key]
                            # 绘制概率密度图
                            a_f.plot_density_with_type(data_dict, title, key + '_' + self.name, self.name)

                        # 特征字段值为数值型
                        else:
                            # 获取数据
                            new_df = pd.DataFrame()
                            new_df['num1'] = self.data[self.name]
                            new_df['num2'] = self.data[key]
                            new_df = new_df.dropna(axis=0, how='any')
                            data1 = new_df['num1'].values
                            data2 = new_df['num2'].values
                            # 获取标题
                            if key not in self.miss:
                                title = "特征%s：散点" % key
                            else:
                                title = "特征%s：散点" % key + ' %s' % self.miss[key]
                            # 绘制散点图
                            a_f.plot_scatter(data1, data2, title, key + '_' + self.name, self.name, key)

                        # 输出数据报告
                        DOC.add_paragraph('关系: ' + key + ' VS ' + self.name, style='List Bullet')
                        DOC.add_picture(r'%s/%s_%s.png' % (self.path, key, self.name), width=Inches(5.8))
        return print('四、每个特征字段与目标字段之间的关系：完毕')

    def relation_multi_feature_with_target(self):
        """
        绘制由2个特征的组合构成的特征和目标特征之间的关系图。
        目标特征是类别特征，
             如果特征组合中全是类别型特征，则绘制分类饼图，
             如果只有一个特征是类别型特征，则绘制箱图
             如果均是连续型特征，则绘制分类散点图，不同类别用颜色区分
        目标特征是连读特征
            如果特征组合中全是连续特征，则绘制三维散点图
            如果只有一个值类别型特征，则绘制分类散点图。
            如果均是类别型特征，则每个特征值得组合绘制箱图
        :return: 图
        """
        DOC.add_heading('五、两个类别特征与目标特征之间的分布关系', level=1)
        # 首先获取
        all_column = self.data.keys()
        # 去除掉一个的，类别型特征去掉含有多个值的
        rest_column = [d for d in all_column if d not in self.one and d not in self.special and d != self.name]
        # 开始绘制图
        for a in range(len(rest_column) - 1):
            for b in range(a + 1, len(rest_column)):
                h, s = rest_column[a], rest_column[b]
                # 类别型缺失值的数理：缺失值变为一类；  对于连续型的特征，缺失值直接去掉。
                new_df = pd.DataFrame()
                new_df['num'] = self.data[self.name]
                new_df['h'] = self.data[h]
                new_df['s'] = self.data[s]
                # 图片标题
                title = '%s VS  %s 的分布' % (h, s)
                # 图片名称
                fig_name = '%s_%s' % (h, s)
                # 坐标轴标签
                if h in self.miss:
                    labelx = '%s,%s' % (h, self.miss[h])
                else:
                    labelx = h
                if s in self.miss:
                    labely = '%s,%s' % (s, self.miss[s])
                else:
                    labely = s
                # 开始判断
                if type(self.data[h]) == object or h in self.f_list:  # 类别型
                    if type(self.data[s]) == object or s in self.f_list:  # 类别型
                        # 数据值填充
                        new_df = new_df.fillna(self.miss_name)
                    else:
                        new_df = new_df.dropna(subset=["s"])
                        new_df = new_df.fillna(self.miss_name)
                else:
                    if type(self.data[s]) == object or s in self.f_list:  # 类别型
                        new_df = new_df.dropna(subset=["h"])
                        new_df = new_df.fillna(self.miss_name)
                    else:
                        new_df = new_df.dropna(axis=0, how='any')
                # 获取数据
                data1 = new_df['h'].values
                data2 = new_df['s'].values
                data3 = new_df['num'].values
                # 目标特征是类别型特征
                if self.data[self.name].dtype == object or self.name in self.f_list:
                    # 如果h，s均是类别型特征
                    if (self.data[h].dtype == object or h in self.f_list) and (self.data[s].dtype == object or s in self.f_list):
                        # 绘制各个特征值组合中的类别的饼图
                        a_f.plot_two_type_type_pie(data1, data2, data3, title, fig_name)
                    # h不是类别型特征
                    elif self.data[h].dtype != object and h not in self.f_list:
                        if self.data[s].dtype == object or s in self.f_list:
                            # 箱图
                            a_f.plot_one_type_type_box(data1, data2, data3, title, fig_name, labelx, labely, self.name)
                        else:
                            # 特征均是连续的
                            a_f.plot_scatter_type(data1, data2, data3, title, fig_name, labelx, labely)
                    # s不是类别型特征
                    else:
                        if self.data[h].dtype == object or h in self.f_list:
                            # 箱图
                            a_f.plot_one_type_type_box(data2, data1, data3, title, fig_name, labely, labelx, self.name)
                        else:
                            # 特征均是连续的
                            a_f.plot_scatter_type(data1, data2, data3, title, fig_name, labelx, labely)
                # 目标特征是连续型特征
                else:
                    # 如果h，s均是类别型特征
                    if (self.data[h].dtype == object or h in self.f_list) and (self.data[s].dtype == object or s in self.f_list):
                        # 特征值组合箱图
                        a_f.plot_two_type_num_box(data1, data2, data3, title, fig_name)
                    # h不是类别型特征
                    elif self.data[h].dtype != object and h not in self.f_list:
                        if self.data[s].dtype == object or s in self.f_list:
                            # 散点
                            a_f.plot_scatter_type(data1, data3, data2, title, fig_name, labelx, self.name)
                        else:
                            # 3d散点图
                            a_f.plot_scatter_3d(data1, data2, data3, title, fig_name, labelx, labely,self.name)
                    # s不是类别型特征
                    else:
                        if self.data[h].dtype == object or h in self.f_list:
                            # 散点
                            a_f.plot_scatter_type(data2, data3, data1, title, fig_name, labely, self.name)
                        else:
                            # 3d散点图
                            a_f.plot_scatter_3d(data1, data2, data3, title, fig_name, labelx, labely, self.name)

        return print('五、组合特征与目标特征的分布：完毕')

    def report_other(self):
        DOC.add_heading('六、其他', level=1)
        DOC.add_paragraph('类别型特征中，不同值较多的字段：')
        for j in self.special:
            DOC.add_paragraph('字段%s，不同值的个数：%d' % (j, len(set(self.data[j].values))), style='List Bullet')
        if self.one:
            DOC.add_paragraph('类别型特征中，只有一个值的字段：')
            for h in self.one:
                DOC.add_paragraph('字段%s，只有一个值' % h, style='List Bullet')

        return print('六、其他：完毕')


# 主函数
if __name__ == "__main__":

    data_report = VIEW()
    data_report.relation_feature_with_target()
    data_report.relation_multi_feature_with_target()
    data_report.report_other()
    DOC.save(r'%s/%s数据报告.docx' % (data_report.path, data_report.report_name))
    print('数据报告生成完毕')


